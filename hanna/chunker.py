from .credentials import ClientCredentials
# from langchain.text_splitter import RecursiveCharacterTextSplitter
from llama_index.core.node_parser import SentenceSplitter
from llama_index.core import Document
from pypdf import PdfReader
import docx
import io


class ChunkText(ClientCredentials):
    def __init__(self, size: int = 300):
        super(ChunkText).__init__()
        self.dir = "./_tmp/"

        self.__rcts = SentenceSplitter(chunk_size=size)

        # RecursiveCharacterTextSplitter(chunk_overlap=0, chunk_size=size, length_function=len)

    def chunk_document(self, path: str):
        corpus = ""
        file = path.split('.')

        if file[-1] == "docx":
            document = docx.Document(self.dir + path)
            print(f"\nReading DOCX [{file}...")

            for x, page in enumerate(document.paragraphs):
                # print(f"PAGE {x + 1}")
                corpus += str(page.text).lower() + " "

        elif file[-1] == "txt":
            file_txt = io.open(self.dir + path, encoding="utf-8")
            corpus = file_txt.read().lower()
            print(f"\nReading TXT [{file_txt.name}]...")
            file_txt.close()

        elif file[-1] == "pdf":
            reader = PdfReader(self.dir + path)
            for page in reader.pages:
                corpus += page.extract_text()
        else:
            return ""

        split_text = self.__rcts.split_text(corpus)

        return split_text

    def chunk_corpus(self, text: str):

        # split_text = self.__rcts.split_text(text)
        chunks = self.__rcts.get_nodes_from_documents([Document(text=text)])
        nodes = [chunk.text for chunk in chunks]

        return nodes